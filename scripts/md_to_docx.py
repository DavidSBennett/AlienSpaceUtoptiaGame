"""
Minimal Markdown -> .docx converter tailored to docs/RULES.md.

Handles the subset of Markdown this rulebook uses: ATX headings (#..####),
horizontal rules (---), blockquotes (>), unordered lists (- ), GitHub tables
(| a | b |), fenced code blocks (```), and inline **bold** / *italic* / `code`.
Not a general Markdown engine — just enough to render our doc cleanly.
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_runs(paragraph, text):
    """Render inline **bold**, *italic*, and `code` into runs."""
    # Split on the inline markers, keeping the delimiters.
    pattern = r'(\*\*.+?\*\*|\*.+?\*|`.+?`)'
    for part in re.split(pattern, text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x30)
        else:
            paragraph.add_run(part)


def split_table_row(line):
    cells = line.strip().strip('|').split('|')
    return [c.strip() for c in cells]


def main(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing fence
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            continue

        # Blank line
        if stripped == '':
            i += 1
            continue

        # Horizontal rule -> a thin spacer / page-ish divider
        if re.fullmatch(r'-{3,}', stripped):
            p = doc.add_paragraph()
            p.add_run('_' * 40).font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,4})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # Strip surrounding emphasis from heading text but keep words.
            heading = doc.add_heading(level=min(level, 4))
            add_runs(heading, text)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            quote_lines = []
            while i < n and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip().lstrip('>').strip())
                i += 1
            p = doc.add_paragraph(style='Intense Quote')
            add_runs(p, ' '.join(quote_lines))
            continue

        # Table (a header row followed by a |---|---| separator)
        if stripped.startswith('|') and i + 1 < n and re.match(r'^\|[\s:\-|]+\|?$', lines[i + 1].strip()):
            header = split_table_row(lines[i])
            i += 2  # skip header + separator
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                rows.append(split_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = 'Light Grid Accent 1'
            for idx, cell_text in enumerate(header):
                cell = table.rows[0].cells[idx]
                cell.paragraphs[0].text = ''
                run = cell.paragraphs[0].add_run('')
                add_runs(cell.paragraphs[0], cell_text)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            for row in rows:
                cells = table.add_row().cells
                for idx in range(len(header)):
                    txt = row[idx] if idx < len(row) else ''
                    cells[idx].paragraphs[0].text = ''
                    add_runs(cells[idx].paragraphs[0], txt)
            continue

        # Unordered list
        if re.match(r'^[-*]\s+', stripped):
            while i < n and re.match(r'^\s*[-*]\s+', lines[i]):
                indent = len(lines[i]) - len(lines[i].lstrip())
                text = re.sub(r'^\s*[-*]\s+', '', lines[i])
                style = 'List Bullet 2' if indent >= 2 else 'List Bullet'
                p = doc.add_paragraph(style=style)
                add_runs(p, text)
                i += 1
            continue

        # Ordered list
        if re.match(r'^\d+\.\s+', stripped):
            while i < n and re.match(r'^\s*\d+\.\s+', lines[i]):
                text = re.sub(r'^\s*\d+\.\s+', '', lines[i])
                p = doc.add_paragraph(style='List Number')
                add_runs(p, text)
                i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f'Wrote {docx_path}')


if __name__ == '__main__':
    main(sys.argv[1], sys.argv[2])
